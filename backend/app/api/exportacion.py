from fastapi import APIRouter, Depends, HTTPException
from fastapi.responses import StreamingResponse
from sqlalchemy.orm import Session
from sqlalchemy import or_
from io import BytesIO
import csv
import io
from app.core.database import get_db
from app.core.security import get_current_user
from app.models.user import User
from app.models.registro import Registro
from app.schemas.schemas import ExportRequest

router = APIRouter(prefix="/api/exportacion", tags=["Exportación"])

FIELD_ORDER = [
    "id", "fuente", "fecha", "link", "titulo", "que", "quien",
    "porque", "datos", "tags", "sector", "orbita", "genero",
    "ambito", "region"
]

FIELD_LABELS = {
    "id": "ID", "fuente": "Fuente", "fecha": "Fecha", "link": "Link",
    "titulo": "Título", "que": "QUÉ", "quien": "QUIÉN",
    "porque": "POR QUÉ", "datos": "DATOS", "tags": "Tags",
    "sector": "Sector", "orbita": "Órbita", "genero": "Género",
    "ambito": "Ámbito", "region": "Región"
}

TRACEABILITY_FIELDS = [
    "titulo_origen", "que_origen", "quien_origen", "porque_origen",
    "datos_origen", "tags_origen", "sector_origen", "orbita_origen"
]


def _get_filtered_registros(data: ExportRequest, db: Session):
    """Get registros based on export request filters."""
    if data.ids:
        return db.query(Registro).filter(Registro.id.in_(data.ids)).all()

    query = db.query(Registro).filter(Registro.estado == "aprobado")
    if data.fuente:
        query = query.filter(Registro.fuente.ilike(f"%{data.fuente}%"))
    if data.sector:
        query = query.filter(Registro.sector == data.sector)
    if data.orbita:
        query = query.filter(Registro.orbita == data.orbita)
    if data.genero:
        query = query.filter(Registro.genero == data.genero)
    if data.ambito:
        query = query.filter(Registro.ambito == data.ambito)
    if data.fecha_desde:
        query = query.filter(Registro.fecha >= data.fecha_desde)
    if data.fecha_hasta:
        query = query.filter(Registro.fecha <= data.fecha_hasta)
    if data.busqueda:
        s = f"%{data.busqueda}%"
        query = query.filter(or_(
            Registro.titulo.ilike(s), Registro.que.ilike(s),
            Registro.quien.ilike(s), Registro.tags.ilike(s)
        ))
    return query.order_by(Registro.fecha.desc()).all()


@router.post("/")
def export_registros(
    data: ExportRequest,
    db: Session = Depends(get_db),
    current_user: User = Depends(get_current_user)
):
    registros = _get_filtered_registros(data, db)

    if not registros:
        raise HTTPException(status_code=404, detail="No hay registros para exportar")

    if data.formato == "csv":
        return _export_csv(registros, data)
    elif data.formato == "xlsx":
        return _export_xlsx(registros, data)
    elif data.formato == "docx":
        return _export_docx(registros, data)
    else:
        raise HTTPException(status_code=400, detail="Formato no soportado")


def _export_csv(registros, data: ExportRequest):
    output = io.StringIO()
    fields = ["titulo"] if data.solo_titulos else FIELD_ORDER
    if data.incluir_trazabilidad and not data.solo_titulos:
        fields = fields + TRACEABILITY_FIELDS

    writer = csv.writer(output)
    writer.writerow([FIELD_LABELS.get(f, f) for f in fields])

    for r in registros:
        row = []
        for f in fields:
            val = getattr(r, f, "")
            if f == "fecha" and val:
                val = val.strftime("%d/%m/%Y") if hasattr(val, 'strftime') else str(val)
            row.append(str(val) if val else "")
        writer.writerow(row)

    output.seek(0)
    return StreamingResponse(
        io.BytesIO(output.getvalue().encode('utf-8-sig')),
        media_type="text/csv",
        headers={"Content-Disposition": "attachment; filename=registros.csv"}
    )


def _export_xlsx(registros, data: ExportRequest):
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side

    wb = Workbook()
    ws = wb.active
    ws.title = "Registros"

    fields = ["titulo"] if data.solo_titulos else FIELD_ORDER
    if data.incluir_trazabilidad and not data.solo_titulos:
        fields = fields + TRACEABILITY_FIELDS

    # Header styling
    header_fill = PatternFill(start_color="2D3748", end_color="2D3748", fill_type="solid")
    header_font = Font(name="Calibri", size=11, bold=True, color="FFFFFF")
    thin_border = Border(
        left=Side(style='thin'), right=Side(style='thin'),
        top=Side(style='thin'), bottom=Side(style='thin')
    )

    for col, field in enumerate(fields, 1):
        cell = ws.cell(row=1, column=col, value=FIELD_LABELS.get(field, field))
        cell.fill = header_fill
        cell.font = header_font
        cell.alignment = Alignment(horizontal="center")
        cell.border = thin_border

    for row_idx, r in enumerate(registros, 2):
        for col, field in enumerate(fields, 1):
            val = getattr(r, field, "")
            if field == "fecha" and val:
                val = val.strftime("%d/%m/%Y") if hasattr(val, 'strftime') else str(val)
            cell = ws.cell(row=row_idx, column=col, value=str(val) if val else "")
            cell.border = thin_border
            cell.alignment = Alignment(wrap_text=True, vertical="top")

    # Auto-width columns
    for col in ws.columns:
        max_length = max(len(str(cell.value or "")) for cell in col)
        ws.column_dimensions[col[0].column_letter].width = min(max_length + 2, 50)

    buf = BytesIO()
    wb.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
        headers={"Content-Disposition": "attachment; filename=registros.xlsx"}
    )


def _export_docx(registros, data: ExportRequest):
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()

    # Title
    title = doc.add_heading("Registros de Noticias", level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Total de registros: {len(registros)}")
    doc.add_paragraph("")

    for r in registros:
        if data.solo_titulos:
            p = doc.add_paragraph()
            p.add_run(f"• {r.titulo or 'Sin título'}").font.size = Pt(11)
            continue

        # Record header
        heading = doc.add_heading(r.titulo or "Sin título", level=2)

        # Metadata
        meta = doc.add_paragraph()
        meta.add_run(f"Fuente: ").bold = True
        meta.add_run(f"{r.fuente}  |  ")
        meta.add_run(f"Fecha: ").bold = True
        fecha_str = r.fecha.strftime("%d/%m/%Y") if r.fecha and hasattr(r.fecha, 'strftime') else str(r.fecha or "")
        meta.add_run(f"{fecha_str}  |  ")
        meta.add_run(f"Sector: ").bold = True
        meta.add_run(f"{r.sector or '-'}  |  ")
        meta.add_run(f"Órbita: ").bold = True
        meta.add_run(f"{r.orbita or '-'}")

        # Content fields
        for field, label in [("que", "QUÉ"), ("quien", "QUIÉN"), ("porque", "POR QUÉ"), ("datos", "DATOS")]:
            val = getattr(r, field, None)
            if val:
                p = doc.add_paragraph()
                run = p.add_run(f"{label}: ")
                run.bold = True
                p.add_run(val)

                if data.incluir_trazabilidad:
                    origen = getattr(r, f"{field}_origen", "ia")
                    origin_run = p.add_run(f"  [{origen.upper()}]")
                    origin_run.font.color.rgb = RGBColor(0x66, 0x66, 0xF1) if origen == "ia" else RGBColor(0x22, 0xC5, 0x5E)
                    origin_run.font.size = Pt(8)

        # Tags
        if r.tags:
            p = doc.add_paragraph()
            p.add_run("Tags: ").bold = True
            p.add_run(r.tags)

        # Link
        p = doc.add_paragraph()
        p.add_run("Link: ").bold = True
        p.add_run(r.link)

        doc.add_paragraph("─" * 60)

    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)

    return StreamingResponse(
        buf,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": "attachment; filename=registros.docx"}
    )
